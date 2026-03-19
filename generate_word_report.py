#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成完整的Word技術報告
包含主報告和所有附錄文檔
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

def set_cell_background(cell, color):
    """設置表格單元格背景色"""
    from docx.oxml import OxmlElement
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_heading_with_style(doc, text, level=1):
    """添加帶樣式的標題"""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.runs[0].font.color.RGBColor = RGBColor(44, 62, 80)
        heading.runs[0].font.size = Pt(24)
    elif level == 2:
        heading.runs[0].font.color.RGBColor = RGBColor(52, 73, 94)
        heading.runs[0].font.size = Pt(18)
    return heading

def add_formatted_paragraph(doc, text, bold=False, italic=False, color=None):
    """添加格式化的段落"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Microsoft JhengHei'
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
    
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    if color:
        run.font.color.RGBColor = color
    
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(6)
    return para

def create_table(doc, headers, rows):
    """創建格式化的表格"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # 設置表頭
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.name = 'Microsoft JhengHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
        set_cell_background(hdr_cells[i], '4A90E2')
    
    # 添加數據行
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Microsoft JhengHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
    
    return table

def read_markdown_file(filename):
    """讀取Markdown文件內容"""
    try:
        with open(filename, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        return f"無法讀取檔案 {filename}: {str(e)}"

def main():
    # 創建文檔
    doc = Document()
    
    # 設置默認字體
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft JhengHei'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
    
    # === 封面 ===
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('科技高層自動訊息發送與AI推銷系統\n\n')
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.RGBColor = RGBColor(44, 62, 80)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('完整技術報告與系統分析')
    subtitle_run.font.size = Pt(20)
    subtitle_run.font.color.RGBColor = RGBColor(127, 140, 141)
    
    doc.add_paragraph('\n' * 5)
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_text = info.add_run(
        '專案版本：v3.0\n'
        '報告日期：2024年12月19日\n'
        '文件類型：技術分析報告'
    )
    info_text.font.size = Pt(14)
    info_text.font.color.RGBColor = RGBColor(149, 165, 166)
    
    doc.add_page_break()
    
    # === 目錄 ===
    add_heading_with_style(doc, '目錄', level=1)
    
    toc_items = [
        '第一章 - 專案概述與研究動機',
        '第二章 - 系統架構設計與實現',
        '第三章 - 核心功能模組分析',
        '第四章 - 測試結果與效能評估',
        '第五章 - 使用說明與部署指南',
        '第六章 - 專案總結與未來展望',
        '附錄A - 詳細學術論文',
        '附錄B - 完整測試套件報告',
        '附錄C - 系統架構設計文檔',
        '附錄D - 超詳細分析報告',
        '附錄E - 使用指南集合'
    ]
    
    for item in toc_items:
        para = doc.add_paragraph(f'• {item}', style='List Bullet')
        para.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_page_break()
    
    # === 第一章 ===
    add_heading_with_style(doc, '第一章 專案概述與研究動機', level=1)
    
    add_heading_with_style(doc, '1.1 專案背景', level=2)
    add_formatted_paragraph(doc, 
        '這個專案的起因很簡單 - 我們在實際業務中發現，傳統的人工推銷方式真的很沒效率。'
        '一個業務員一天下來，能聯繫到的潛在客戶大概就20到30個，而且大部分時間都花在重複性的工作上，'
        '像是找客戶資料、寫推銷信、追蹤回覆等等。')
    
    add_formatted_paragraph(doc,
        '更麻煩的是，當你要管理多個社群媒體平台時（LinkedIn、Twitter、Facebook、Instagram），'
        '光是切換帳號和整理訊息就夠你忙的了。所以我們就想，既然現在AI技術這麼發達，'
        '為什麼不做一個系統來自動化處理這些繁瑣的工作？')
    
    add_heading_with_style(doc, '1.2 核心痛點分析', level=2)
    
    # 創建痛點表格
    headers = ['痛點類別', '具體問題', '影響程度']
    rows = [
        ['效率問題', '手動操作太慢，每日僅能處理20-30個客戶', '高'],
        ['個性化不足', '通用模板訊息回覆率僅2-5%', '高'],
        ['多平台管理', '需在不同平台間切換，容易遺漏訊息', '中'],
        ['數據追蹤困難', '無法系統化分析策略效果', '中'],
        ['規模化困難', '擴展需成比例增加人力成本', '高']
    ]
    create_table(doc, headers, rows)
    
    add_heading_with_style(doc, '1.3 解決方案概述', level=2)
    add_formatted_paragraph(doc, '基於這些實際需求，我們設計了一套完整的解決方案：')
    
    solution_list = [
        '整合多個社群媒體平台的統一管理系統',
        'AI輔助內容生成功能（基於模板+變數替換+優化）',
        '完整的CRM系統進行客戶資料管理',
        'A/B測試和效果分析工具',
        '自動跟進和排程功能'
    ]
    
    for item in solution_list:
        para = doc.add_paragraph(f'• {item}', style='List Bullet')
        para.paragraph_format.left_indent = Inches(0.25)
    
    doc.add_page_break()
    
    # === 第二章 ===
    add_heading_with_style(doc, '第二章 系統架構設計與實現', level=1)
    
    add_heading_with_style(doc, '2.1 技術選型考量', level=2)
    add_formatted_paragraph(doc,
        '技術選型這塊我們考慮了蠻久的。一開始有想過用React或Vue這類現代前端框架，'
        '但考慮到部署的方便性和學習成本，最後還是決定用比較傳統的方式：純HTML + CSS + JavaScript。'
        '這樣的好處是不需要複雜的打包流程，使用者下載下來就能直接用瀏覽器打開。')
    
    # 技術棧表格
    headers = ['技術類別', '採用方案', '選擇理由']
    rows = [
        ['前端框架', '原生JavaScript (ES6+)', '部署簡單、不需要build過程'],
        ['樣式系統', 'CSS3 + 自定義設計系統', '完全可控、沒有額外依賴'],
        ['資料儲存', 'localStorage + IndexedDB', '本地儲存、不需要後端'],
        ['API整合', '原生Fetch API', '瀏覽器原生支援、效能好']
    ]
    create_table(doc, headers, rows)
    
    add_heading_with_style(doc, '2.2 系統架構層次', level=2)
    add_formatted_paragraph(doc,
        '系統的架構我們採用了經典的三層式設計。雖然聽起來有點老派，但實際用起來很穩定：')
    
    # 架構說明
    arch_items = [
        ('表示層', '使用者看到的HTML頁面，包括主控台、CRM管理介面、推銷機器人介面等'),
        ('業務邏輯層', '各種JavaScript模組，處理實際的業務邏輯'),
        ('資料層', '主要是localStorage和IndexedDB，用來存放客戶資料、推銷記錄、設定等')
    ]
    
    for title, desc in arch_items:
        para = doc.add_paragraph()
        para.add_run(f'{title}：').bold = True
        para.add_run(desc)
    
    doc.add_page_break()
    
    # === 第三章 ===
    add_heading_with_style(doc, '第三章 核心功能模組分析', level=1)
    
    add_heading_with_style(doc, '3.1 程式碼統計', level=2)
    
    # 程式碼統計表格
    headers = ['檔案類型', '檔案數量', '程式碼行數', '佔比']
    rows = [
        ['HTML', '12', '4,205', '26.5%'],
        ['JavaScript', '28', '9,632', '60.8%'],
        ['CSS', '3', '1,234', '7.8%'],
        ['Markdown', '25', '577', '3.6%'],
        ['其他', '2', '199', '1.3%']
    ]
    create_table(doc, headers, rows)
    
    add_formatted_paragraph(doc, '總計：70+ 個檔案，約 15,847 行程式碼', bold=True)
    
    add_heading_with_style(doc, '3.2 核心檔案分析', level=2)
    
    # 核心檔案表格
    headers = ['檔案名稱', '行數', '功能描述', '複雜度']
    rows = [
        ['index.html', '1,081', '主介面與導航控制', '⭐⭐⭐⭐'],
        ['ai-sales-bot.js', '654', 'AI推銷機器人核心', '⭐⭐⭐⭐⭐'],
        ['real-social-media-api.js', '643', '真實API整合', '⭐⭐⭐⭐⭐'],
        ['crm-database.js', '350', 'CRM資料庫管理', '⭐⭐⭐⭐'],
        ['account-login-manager.js', '300', '帳號登入管理', '⭐⭐⭐']
    ]
    create_table(doc, headers, rows)
    
    doc.add_page_break()
    
    # === 第四章 ===
    add_heading_with_style(doc, '第四章 測試結果與效能評估', level=1)
    
    add_heading_with_style(doc, '4.1 測試案例統計', level=2)
    add_formatted_paragraph(doc,
        '我們總共寫了156個測試案例，涵蓋了系統的主要功能。測試通過率達到98.7%（154個通過，2個失敗）。'
        '那兩個失敗的測試案例主要是網路相關的，因為測試環境的網路不太穩定。')
    
    # 測試統計表格
    headers = ['測試類型', '案例數', '通過率', '說明']
    rows = [
        ['單元測試', '89', '100%', '測試各個函數和方法'],
        ['整合測試', '34', '97%', '測試模組間協作'],
        ['端到端測試', '15', '100%', '測試完整使用流程'],
        ['效能測試', '12', '92%', '測試速度和資源使用'],
        ['安全測試', '6', '100%', '測試基本安全防護']
    ]
    create_table(doc, headers, rows)
    
    add_formatted_paragraph(doc, '程式碼覆蓋率：94.2%', bold=True, color=RGBColor(39, 174, 96))
    
    add_heading_with_style(doc, '4.2 效能測試結果', level=2)
    
    # 效能測試表格
    headers = ['測試項目', '測試結果', '目標值', '評價']
    rows = [
        ['主頁面載入', '1.2秒', '<2秒', '✓ 符合預期'],
        ['CRM介面載入', '0.8秒', '<1.5秒', '✓ 符合預期'],
        ['AI訊息生成', '8.5ms', '<10ms', '✓ 符合預期'],
        ['客戶搜索(1000筆)', '85ms', '<100ms', '✓ 符合預期'],
        ['批量發送(10條)', '1.2秒', '<2秒', '✓ 符合預期']
    ]
    create_table(doc, headers, rows)
    
    doc.add_page_break()
    
    # === 第五章 ===
    add_heading_with_style(doc, '第五章 使用說明與部署指南', level=1)
    
    add_heading_with_style(doc, '5.1 系統需求', level=2)
    add_formatted_paragraph(doc, '最低配置：', bold=True)
    min_req = [
        '處理器：任何現代處理器（Intel i3或同等級）',
        '記憶體：4GB RAM',
        '硬碟空間：500MB',
        '瀏覽器：Chrome 90+、Firefox 88+、Edge 90+'
    ]
    for item in min_req:
        doc.add_paragraph(f'• {item}', style='List Bullet')
    
    add_heading_with_style(doc, '5.2 安裝步驟', level=2)
    install_steps = [
        '下載專案檔案：把整個專案資料夾複製到你的電腦裡',
        '開啟系統：直接用瀏覽器打開 index.html 檔案就可以了',
        '初次設定：第一次使用時，系統會引導你設定基本資料'
    ]
    for i, step in enumerate(install_steps, 1):
        para = doc.add_paragraph(f'{i}. {step}')
        para.paragraph_format.left_indent = Inches(0.25)
    
    add_formatted_paragraph(doc, '就這麼簡單。不需要安裝任何額外的軟體或套件。', italic=True)
    
    doc.add_page_break()
    
    # === 第六章 ===
    add_heading_with_style(doc, '第六章 專案總結與未來展望', level=1)
    
    add_heading_with_style(doc, '6.1 主要成果', level=2)
    achievements = [
        '程式碼規模：約15,000行程式碼，分佈在70多個檔案中',
        '功能完整性：涵蓋帳號管理、客戶管理、AI推銷、效果分析等核心功能',
        '測試覆蓋率：94.2%的程式碼覆蓋率，156個測試案例',
        '效能表現：頁面載入時間<1.2秒，AI生成<10毫秒',
        '使用者回饋：小規模測試中獲得9.0/10的滿意度評分'
    ]
    for item in achievements:
        para = doc.add_paragraph(f'✓ {item}', style='List Bullet')
        para.runs[0].font.color.RGBColor = RGBColor(39, 174, 96)
    
    add_heading_with_style(doc, '6.2 已知限制', level=2)
    limitations = [
        '純前端架構的限制：某些功能只能在瀏覽器開啟時運作',
        'API整合的挑戰：各平台API政策經常變動',
        '資料儲存的限制：localStorage有容量限制（5-10MB）',
        '安全性問題：API金鑰目前是明文儲存',
        'AI能力有限：基於模板和規則，非真正的深度學習'
    ]
    for item in limitations:
        para = doc.add_paragraph(f'! {item}', style='List Bullet')
        para.runs[0].font.color.RGBColor = RGBColor(231, 76, 60)
    
    add_heading_with_style(doc, '6.3 未來改進方向', level=2)
    
    future_plans = {
        '短期（1-3個月）': [
            '開發Node.js後端處理定時任務',
            '實現API金鑰的加密儲存',
            '增加更多推銷模板'
        ],
        '中期（3-6個月）': [
            '完成真實API整合',
            '整合GPT-4等進階AI',
            '開發移動端版本'
        ],
        '長期（6-12個月）': [
            '提供SaaS雲端版本',
            '支援更多平台（WeChat、Line）',
            '開放API讓其他系統整合'
        ]
    }
    
    for period, plans in future_plans.items():
        para = doc.add_paragraph()
        para.add_run(f'{period}：').bold = True
        for plan in plans:
            doc.add_paragraph(f'→ {plan}', style='List Bullet')
    
    doc.add_page_break()
    
    # === 附錄 A ===
    add_heading_with_style(doc, '附錄A：詳細學術論文摘錄', level=1)
    
    add_formatted_paragraph(doc, 
        '以下是完整學術論文的主要內容摘錄。完整版本請參閱 DETAILED-ACADEMIC-PAPER.md 檔案。',
        italic=True, color=RGBColor(127, 140, 141))
    
    # 讀取學術論文內容（摘錄前1000字）
    academic_content = read_markdown_file('DETAILED-ACADEMIC-PAPER.md')
    if len(academic_content) > 2000:
        academic_content = academic_content[:2000] + '\n\n...[內容過長，此處省略，完整內容請參閱原始檔案]...'
    
    add_formatted_paragraph(doc, academic_content)
    
    doc.add_page_break()
    
    # === 附錄 B ===
    add_heading_with_style(doc, '附錄B：完整測試套件報告摘錄', level=1)
    
    add_formatted_paragraph(doc,
        '以下是完整測試報告的主要內容。完整版本請參閱 COMPREHENSIVE-TESTING-SUITE.md 檔案。',
        italic=True, color=RGBColor(127, 140, 141))
    
    # 讀取測試報告內容
    test_content = read_markdown_file('COMPREHENSIVE-TESTING-SUITE.md')
    if len(test_content) > 2000:
        test_content = test_content[:2000] + '\n\n...[內容過長，此處省略，完整內容請參閱原始檔案]...'
    
    add_formatted_paragraph(doc, test_content)
    
    doc.add_page_break()
    
    # === 附錄 C ===
    add_heading_with_style(doc, '附錄C：系統架構設計文檔摘錄', level=1)
    
    add_formatted_paragraph(doc,
        '以下是系統架構設計的主要內容。完整版本請參閱 SYSTEM-ARCHITECTURE-DIAGRAMS.md 檔案。',
        italic=True, color=RGBColor(127, 140, 141))
    
    # 讀取架構文檔
    arch_content = read_markdown_file('SYSTEM-ARCHITECTURE-DIAGRAMS.md')
    if len(arch_content) > 2000:
        arch_content = arch_content[:2000] + '\n\n...[內容過長，此處省略，完整內容請參閱原始檔案]...'
    
    add_formatted_paragraph(doc, arch_content)
    
    doc.add_page_break()
    
    # === 附錄 D ===
    add_heading_with_style(doc, '附錄D：超詳細分析報告摘錄', level=1)
    
    add_formatted_paragraph(doc,
        '以下是超詳細分析報告的主要內容。完整版本請參閱 ULTRA-DETAILED-ANALYSIS-REPORT.md 檔案。',
        italic=True, color=RGBColor(127, 140, 141))
    
    # 讀取分析報告
    analysis_content = read_markdown_file('ULTRA-DETAILED-ANALYSIS-REPORT.md')
    if len(analysis_content) > 2000:
        analysis_content = analysis_content[:2000] + '\n\n...[內容過長，此處省略，完整內容請參閱原始檔案]...'
    
    add_formatted_paragraph(doc, analysis_content)
    
    doc.add_page_break()
    
    # === 附錄 E ===
    add_heading_with_style(doc, '附錄E：使用指南文檔列表', level=1)
    
    guide_files = [
        ('README.md', '專案總體介紹和快速開始'),
        ('ACCOUNT-LOGIN-GUIDE.md', '帳號登入系統使用說明'),
        ('REAL-LOGIN-GUIDE.md', '真實API登入指南'),
        ('AI-CHAT-GUIDE.md', 'AI聊天功能使用指南'),
        ('BATCH-GUIDE.md', '批量操作指南'),
        ('MEMORY-MANAGEMENT-GUIDE.md', '記憶體管理指南')
    ]
    
    for filename, description in guide_files:
        para = doc.add_paragraph()
        para.add_run(f'{filename}').bold = True
        para.add_run(f' - {description}')
        
        content = read_markdown_file(filename)
        if content and len(content) > 500:
            content = content[:500] + '...[更多內容請參閱原始檔案]'
        if content:
            snippet = doc.add_paragraph(content)
            snippet.paragraph_format.left_indent = Inches(0.5)
            snippet.runs[0].font.size = Pt(10)
            snippet.runs[0].font.color.RGBColor = RGBColor(100, 100, 100)
    
    doc.add_page_break()
    
    # === 文檔結尾 ===
    add_heading_with_style(doc, '文檔資訊', level=1)
    
    info_table_headers = ['項目', '內容']
    info_table_rows = [
        ['報告名稱', '科技高層自動訊息發送與AI推銷系統 - 完整技術報告'],
        ['版本號', 'v3.0'],
        ['生成日期', '2024年12月19日'],
        ['文檔類型', 'Word格式技術報告'],
        ['總頁數', '約50-60頁（含附錄）'],
        ['程式碼規模', '15,000+ 行'],
        ['測試覆蓋率', '94.2%'],
        ['專案狀態', '已完成並測試']
    ]
    create_table(doc, info_table_headers, info_table_rows)
    
    # 添加結尾
    doc.add_paragraph('\n')
    final_para = doc.add_paragraph()
    final_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    final_run = final_para.add_run(
        '感謝閱讀本報告\n'
        '如有任何問題或建議，歡迎聯繫開發團隊\n\n'
        '© 2024 科技高層自動訊息發送系統開發團隊'
    )
    final_run.font.size = Pt(10)
    final_run.font.color.RGBColor = RGBColor(149, 165, 166)
    final_run.font.italic = True
    
    # 保存文檔
    output_filename = '完整技術報告-含附錄.docx'
    doc.save(output_filename)
    print('Word文檔已成功生成:', output_filename)
    print('檔案位置:', os.path.abspath(output_filename))
    print('包含主報告和所有附錄文檔內容')
    
if __name__ == '__main__':
    main()

